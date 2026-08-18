#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Extract text from DOCX file"""

from docx import Document
import sys
import os

docx_path = "docs/00-参考资料/数尔安防官方资料/热成像双光云台相机SR-UPA810T609.docx"

if not os.path.exists(docx_path):
    print(f"文件不存在: {docx_path}")
    sys.exit(1)

try:
    doc = Document(docx_path)
    output_lines = []
    output_lines.append(f"\n{'='*60}")
    output_lines.append(f"文件: 热成像双光云台相机SR-UPA810T609.docx")
    output_lines.append(f"{'='*60}")
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            output_lines.append(para.text)
    
    # Also extract tables
    output_lines.append(f"\n{'='*60}")
    output_lines.append("表格内容:")
    output_lines.append(f"{'='*60}")
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                output_lines.append(' | '.join(cells))
    
    print('\n'.join(output_lines))
except Exception as e:
    print(f"错误: {e}")
    sys.exit(1)
